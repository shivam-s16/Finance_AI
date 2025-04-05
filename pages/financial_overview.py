import os
from crewai import Agent, Task, Crew
from crewai_tools import ScrapeWebsiteTool
from dotenv import load_dotenv
from docx import Document
from io import BytesIO
load_dotenv()
import streamlit as st
# LLM object and API Key


os.environ["SERPER_API_KEY"] = os.getenv("SERPER_API_KEY")
os.environ["GOOGLE_API_KEY"] = os.getenv("GOOGLE_API_KEY")
from docx import Document
from io import BytesIO
import base64
from langchain_google_genai import ChatGoogleGenerativeAI
from dotenv import load_dotenv
from crewai_tools import (
    SerperDevTool,
    WebsiteSearchTool
)

scrape_tool = ScrapeWebsiteTool()
search_tool = SerperDevTool()
# Get BRAVE_API_KEY from environment variables
# api_key = os.getenv('BRAVE_API_KEY')

#search_tool = BraveSearch.from_api_key(api_key=api_key,
#                                       search_kwargs={"count": 3})
#openbb_tool = OpenBBTools()
def generate_docx(result):
    doc = Document()
    doc.add_heading('Fin advice', 0)
    doc.add_paragraph(result)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
#
llm=ChatGoogleGenerativeAI(model="gemini-1.5-flash",
                           verbose=True,
                           temperature=0.7,
                           GOOGLE_API_KEY="")


researcher = Agent(
    role='Senior Research Analyst',
    goal='Uncover insights into the company ({company})',
    backstory=
    """You work as a research analyst at Goldman Sachs, focusing on fundamental 
    research for tech companies""",
    verbose=True,
    allow_delegation=False,
    tools=[search_tool,scrape_tool],
    llm=llm)

visionary = Agent(
    role='Visionary',
    goal='Deep thinking on the implications of an analysis',
    backstory=
    """you are a visionary technologist with a keen eye for identifying emerging 
    trends and predicting their potential impact on various industries. Your ability 
    to think critically and connect seemingly disparate dots allows you to anticipate 
    disruptive technologies and their far-reaching implications.""",
    verbose=True,
    allow_delegation=False,
    tools=[ search_tool,scrape_tool],
    llm=llm)

writer = Agent(
    role='Senior editor',
    goal='Writes professional quality articles that are easy to understand',
    backstory=
    """You are a details-oriented senior editor at the Wall Street Journal 
    known for your insightful and engaging articles. You transform complex 
    concepts into factual and impactful narratives.""",
    verbose=True,
    llm=llm,
    tools=[search_tool,scrape_tool],
    allow_delegation=True)

# Create tasks for your agents
task1 = Task(
    description=
    """please conduct a comprehensive analysis of ({company})'s latest SEC 10-K filing. The analysis should include the following key points:

  Business Overview: Briefly describe ({company})'s business model, its products and services, and its target market.

  Risk Factors: Identify and discuss the major risk factors that ({company}) has disclosed in its 10-K filing.

  Management's Discussion and Analysis (MD&A): Summarize the key points from the MD&A section, including any significant changes in operations, financial condition, or liquidity.

  Competitive Landscape: Discuss ({company})'s competitive position in its industry and how it compares to its major competitors.

  Future Outlook: Based on the information in the 10-K filing and your analysis, provide a brief outlook on ({company})'s future performance.

  Please ensure that all information is sourced from ({company})'s latest SEC 10-K filing and that the analysis is unbiased and factual.""",
    expected_output="Full analysis report in bullet points",
    agent=researcher)

task2 = Task(
    description=
    """Using the insights provided by the Senior Research Analyst, think through deeply the future implications of the points that are made. Consider the following questions as you craft your response:

What are the current limitations or pain points that this technology mentioned in the Senior Research Analyst report could address?

How might this technology disrupt traditional business models and create new opportunities for innovation?

What are the potential risks and challenges associated with the adoption of this technology, and how might they be mitigated?

How could this technology impact consumers, employees, and society as a whole?

What are the long-term implications of this technology, and how might it shape the future of the industry?

Provide a detailed analysis of the technology's potential impact, backed by relevant examples, data, and insights. Your response should demonstrate your ability to think strategically, anticipate future trends, and articulate complex ideas in a clear and compelling manner.""",
    expected_output="Analysis report with deeper insights in implications",
    agent=visionary)

task3 = Task(
    description=
    """Using the insights provided by the Senior Research Analyst and Visionary,please craft an expertly styled report that is targeted towards the investor community. Make sure to also include the long-term implications insights that your co-worker, Visionary, has shared.  
    
    Please ensure that the report is written in a professional tone and style, and that all information is sourced from ({company})'s latest SEC 10-K filing. Write in a format and style worthy to be published in the wall street journal.""",
    expected_output=
    "A detailed comprehensive report on NVDIA that expertly presents the research done by your co-worker, Senior Research Analyst and Visionary",
    agent=writer)
def app():
# Instantiate your crew with a sequential process
    crew = Crew(
        agents=[researcher, visionary, writer],
        tasks=[task1, task2, task3],
        verbose=2,  # You can set it to 1 or 2 to different logging levels
    )
    
    company = st.text_input("Stock Selection")
    financial_inputs=({
        'company': company
    })
    # Get your crew to work!
    result = crew.kickoff(inputs=financial_inputs)


    docx_file = generate_docx(result)
